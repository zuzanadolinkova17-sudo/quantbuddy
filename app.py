# app.py — QuantBuddy AI Asistent
# Inteligentní průvodce kvantitativní analýzou pro studenty
# Spuštění: streamlit run app.py

import io
import tempfile
from datetime import datetime
import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from scipy import stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========================================
# KONFIGURACE
# ========================================

st.set_page_config(
    page_title="QuantBuddy AI — Tvůj statistický asistent",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========================================
# CUSTOM CSS — Moderní design
# ========================================

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    
    * {
        font-family: 'Inter', sans-serif;
    }
    
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 3rem 2rem;
        border-radius: 20px;
        text-align: center;
        color: white;
        margin-bottom: 2rem;
        box-shadow: 0 10px 40px rgba(102, 126, 234, 0.3);
    }
    
    .main-header h1 {
        font-size: 3rem;
        font-weight: 700;
        margin: 0;
    }
    
    .main-header p {
        font-size: 1.3rem;
        margin: 0.5rem 0 0 0;
        opacity: 0.95;
    }
    
    .step-card {
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
        padding: 2rem;
        border-radius: 15px;
        border-left: 5px solid #667eea;
        margin: 1.5rem 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .info-box {
        background: #e3f2fd;
        border-left: 4px solid #2196F3;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    
    .warning-box {
        background: #fff3e0;
        border-left: 4px solid #ff9800;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    
    .success-box {
        background: #e8f5e9;
        border-left: 4px solid #4caf50;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    
    .result-card {
        background: white;
        padding: 2rem;
        border-radius: 15px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        margin: 1.5rem 0;
        border-top: 4px solid #667eea;
    }
    
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
    }
    
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        border: none;
        font-size: 1.1rem;
        transition: all 0.3s;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);
    }
    
    .stButton>button:hover {
        transform: translateY(-3px);
        box-shadow: 0 6px 25px rgba(102, 126, 234, 0.5);
    }
    
    .chat-bubble {
        background: #f5f5f5;
        padding: 1.5rem;
        border-radius: 15px;
        margin: 1rem 0;
        position: relative;
    }
    
    .chat-bubble::before {
        content: "🤖";
        position: absolute;
        left: -40px;
        font-size: 2rem;
    }
    
    .progress-container {
        background: #f0f0f0;
        border-radius: 10px;
        padding: 0.5rem;
        margin: 2rem 0;
    }
    
    .tooltip-text {
        font-size: 0.9rem;
        color: #666;
        font-style: italic;
    }
</style>
""", unsafe_allow_html=True)

# ========================================
# SESSION STATE
# ========================================

if 'step' not in st.session_state:
    st.session_state.step = 1
if 'df' not in st.session_state:
    st.session_state.df = None
if 'research_goal' not in st.session_state:
    st.session_state.research_goal = None
if 'recommended_analysis' not in st.session_state:
    st.session_state.recommended_analysis = None

# ========================================
# POMOCNÉ FUNKCE
# ========================================

def detect_var_types(df: pd.DataFrame, cat_threshold: int = 10):
    """Inteligentní detekce typů proměnných."""
    types = {}
    descriptions = {}
    for col in df.columns:
        s = df[col]
        if pd.api.types.is_numeric_dtype(s):
            nunq = s.dropna().nunique()
            if nunq <= cat_threshold:
                types[col] = "kategorická"
                descriptions[col] = f"Číselná, ale jen {nunq} různých hodnot → považuji za kategorickou"
            else:
                types[col] = "numerická"
                descriptions[col] = f"Číselná s {nunq} unikátními hodnotami"
        else:
            types[col] = "kategorická"
            nunq = s.dropna().nunique()
            descriptions[col] = f"Textová s {nunq} kategoriemi"
    return types, descriptions

def recommend_analysis(df, types, goal):
    """AI doporučení analýzy na základě dat a cíle."""
    num_cols = [c for c, t in types.items() if t == "numerická"]
    cat_cols = [c for c, t in types.items() if t == "kategorická"]
    
    recommendations = []
    
    if goal == "vztah":
        if len(num_cols) >= 2:
            recommendations.append({
                "name": "Korelace (Pearsonova/Spearmanova)",
                "icon": "🔗",
                "reason": f"Máš {len(num_cols)} numerických proměnných. Korelace ukáže, jak spolu souvisí.",
                "suitable_for": "Zjištění síly a směru vztahu mezi dvěma číselnými proměnnými",
                "example": "Např. vztah mezi studijním časem a známkami",
                "variables": {"x": num_cols, "y": num_cols},
                "type": "correlation"
            })
            recommendations.append({
                "name": "Lineární regrese",
                "icon": "📈",
                "reason": "Můžeš předpovědět jednu proměnnou z druhé.",
                "suitable_for": "Predikce hodnoty jedné proměnné na základě druhé",
                "example": "Např. předpověď známky podle studijního času",
                "variables": {"x": num_cols, "y": num_cols},
                "type": "regression"
            })
        if len(cat_cols) >= 2:
            recommendations.append({
                "name": "Chí-kvadrát test",
                "icon": "🎲",
                "reason": f"Máš {len(cat_cols)} kategorických proměnných. Chí-kvadrát ukáže, zda spolu souvisí.",
                "suitable_for": "Zjištění asociace mezi kategoriemi",
                "example": "Např. souvislost mezi pohlavím a oborem",
                "variables": {"var1": cat_cols, "var2": cat_cols},
                "type": "chi2"
            })
    
    elif goal == "porovnání":
        if len(cat_cols) >= 1 and len(num_cols) >= 1:
            # Zkontroluj, jestli nějaká kategorická má přesně 2 úrovně
            two_level_cats = []
            for col in cat_cols:
                if df[col].nunique() == 2:
                    two_level_cats.append(col)
            
            if two_level_cats:
                recommendations.append({
                    "name": "T-test (porovnání 2 skupin)",
                    "icon": "⚖️",
                    "reason": f"Máš kategorickou proměnnou se 2 skupinami a numerickou proměnnou.",
                    "suitable_for": "Porovnání průměrů mezi dvěma skupinami",
                    "example": "Např. porovnání platu mužů vs. žen",
                    "variables": {"group": two_level_cats, "outcome": num_cols},
                    "type": "ttest"
                })
            
            multi_level_cats = [c for c in cat_cols if df[c].nunique() > 2]
            if multi_level_cats:
                recommendations.append({
                    "name": "ANOVA (porovnání více skupin)",
                    "icon": "📊",
                    "reason": f"Máš kategorickou proměnnou s více než 2 skupinami.",
                    "suitable_for": "Porovnání průměrů mezi 3+ skupinami",
                    "example": "Např. porovnání platů mezi obory",
                    "variables": {"group": multi_level_cats, "outcome": num_cols},
                    "type": "anova"
                })
    
    elif goal == "popis":
        recommendations.append({
            "name": "Deskriptivní statistika",
            "icon": "📋",
            "reason": "Ukážu ti základní charakteristiky tvých dat.",
            "suitable_for": "Popis dat (průměr, medián, rozptyl, četnosti)",
            "example": "Např. průměrný věk, nejčastější odpověď",
            "variables": {"cols": list(df.columns)},
            "type": "descriptive"
        })
    
    if not recommendations:
        recommendations.append({
            "name": "Nejsem si jistý",
            "icon": "🤔",
            "reason": "Tvoje data nebo cíl nejsou jednoznačné. Zkus mi říct víc!",
            "suitable_for": "",
            "example": "",
            "variables": {},
            "type": None
        })
    
    return recommendations

def check_assumptions(df, analysis_type, var1, var2=None):
    """Kontrola statistických předpokladů."""
    warnings = []
    tips = []
    
    if analysis_type in ["correlation", "regression", "ttest", "anova"]:
        # Kontrola velikosti vzorku
        n = len(df[[var1, var2]].dropna()) if var2 else len(df[var1].dropna())
        if n < 30:
            warnings.append(f"⚠️ Máš jen {n} pozorování. Ideálně alespoň 30 pro spolehlivé výsledky.")
        else:
            tips.append(f"✅ Velikost vzorku ({n}) je v pořádku!")
        
        # Kontrola normality pro numerické proměnné
        if analysis_type in ["correlation", "regression", "ttest"]:
            for var in [var1, var2] if var2 else [var1]:
                if pd.api.types.is_numeric_dtype(df[var]):
                    clean_data = df[var].dropna()
                    if len(clean_data) >= 3:
                        _, p = stats.shapiro(clean_data[:5000])  # Max 5000 pro Shapiro
                        if p < 0.05:
                            warnings.append(f"⚠️ Proměnná '{var}' není normálně rozložená. Zvaž Spearmanovu korelaci místo Pearsonovy.")
        
        # Kontrola chybějících hodnot
        missing = df[[var1, var2]].isnull().sum().sum() if var2 else df[var1].isnull().sum()
        if missing > 0:
            warnings.append(f"⚠️ {missing} chybějících hodnot bude vyřazeno z analýzy.")
    
    return warnings, tips

# ========================================
# HLAVIČKA
# ========================================

st.markdown("""
<div class="main-header">
    <h1>🤖 QuantBuddy AI</h1>
    <p>Tvůj inteligentní asistent pro kvantitativní analýzu</p>
    <p style="font-size: 1rem; margin-top: 1rem;">Nahraji data → Řeknu ti cíl → Dostanu hotové výsledky pro závěrečnou práci ✨</p>
</div>
""", unsafe_allow_html=True)

# Progress bar
progress = (st.session_state.step - 1) / 3
st.progress(progress)
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown(f"**{'✅' if st.session_state.step > 1 else '1️⃣'} Krok 1: Data**")
with col2:
    st.markdown(f"**{'✅' if st.session_state.step > 2 else '2️⃣' if st.session_state.step == 2 else '⚪'} Krok 2: Cíl**")
with col3:
    st.markdown(f"**{'✅' if st.session_state.step > 3 else '3️⃣' if st.session_state.step == 3 else '⚪'} Krok 3: Analýza**")

st.markdown("---")

# ========================================
# KROK 1: NAHRÁNÍ DAT
# ========================================

if st.session_state.step == 1:
    st.markdown('<div class="step-card">', unsafe_allow_html=True)
    st.markdown("## 📂 Krok 1: Nahraj svá data")
    st.markdown("Přetáhni sem svůj Excel nebo CSV soubor. Podívám se na něj a řeknu ti, co v něm vidím.")
    
    file = st.file_uploader("📁 Vyber soubor", type=["csv", "xlsx"], label_visibility="collapsed")
    
    if file:
        try:
            # Načtení souboru
            if file.name.lower().endswith(".xlsx"):
                xls = pd.ExcelFile(file)
                if len(xls.sheet_names) > 1:
                    sheet_name = st.selectbox("📋 Vyber list z Excelu:", xls.sheet_names)
                else:
                    sheet_name = xls.sheet_names[0]
                df = pd.read_excel(file, sheet_name=sheet_name, engine='openpyxl')
            else:
                try:
                    df = pd.read_csv(file, encoding='utf-8')
                except:
                    file.seek(0)
                    df = pd.read_csv(file, encoding='latin1', sep=';')
            
            st.session_state.df = df
            
            # Analýza dat
            types, descriptions = detect_var_types(df)
            
            st.markdown('<div class="success-box">', unsafe_allow_html=True)
            st.markdown(f"### ✅ Super! Tvá data jsou načtena")
            st.markdown(f"**Počet respondentů:** {df.shape[0]}")
            st.markdown(f"**Počet proměnných:** {df.shape[1]}")
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Náhled dat
            with st.expander("👀 Podívej se na data (prvních 10 řádků)"):
                st.dataframe(df.head(10), use_container_width=True)
            
            # Detekce typů
            st.markdown("### 🔍 Co jsem v datech našel:")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**📊 Numerické proměnné** (čísla)")
                num_vars = [k for k, v in types.items() if v == "numerická"]
                if num_vars:
                    for var in num_vars:
                        st.markdown(f"- **{var}** · {descriptions[var]}")
                else:
                    st.markdown("*Žádné nenalezeny*")
            
            with col2:
                st.markdown("**🏷️ Kategorické proměnné** (skupiny)")
                cat_vars = [k for k, v in types.items() if v == "kategorická"]
                if cat_vars:
                    for var in cat_vars:
                        st.markdown(f"- **{var}** · {descriptions[var]}")
                else:
                    st.markdown("*Žádné nenalezeny*")
            
            # Kontrola kvality dat
            st.markdown("### 🩺 Kontrola kvality dat:")
            issues = []
            
            missing_total = df.isnull().sum().sum()
            if missing_total > 0:
                issues.append(f"⚠️ Celkem {missing_total} chybějících hodnot")
            
            duplicates = df.duplicated().sum()
            if duplicates > 0:
                issues.append(f"⚠️ {duplicates} duplicitních řádků")
            
            if issues:
                st.markdown('<div class="warning-box">', unsafe_allow_html=True)
                for issue in issues:
                    st.markdown(issue)
                st.markdown("💡 **Tip:** To je OK, appka s tím umí pracovat. Chybějící hodnoty vyřadím z analýzy.")
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="success-box">', unsafe_allow_html=True)
                st.markdown("✅ Data vypadají skvěle! Žádné chybějící hodnoty ani duplicity.")
                st.markdown('</div>', unsafe_allow_html=True)
            
            if st.button("➡️ Pokračovat na Krok 2", use_container_width=True):
                st.session_state.step = 2
                st.rerun()
        
        except Exception as e:
            st.error(f"❌ Nepodařilo se načíst soubor: {e}")
    
    st.markdown('</div>', unsafe_allow_html=True)

# ========================================
# KROK 2: VÝZKUMNÝ CÍL
# ========================================

elif st.session_state.step == 2:
    df = st.session_state.df
    types, _ = detect_var_types(df)
    
    st.markdown('<div class="step-card">', unsafe_allow_html=True)
    st.markdown("## 🎯 Krok 2: Co chceš zjistit?")
    st.markdown('<div class="chat-bubble">', unsafe_allow_html=True)
    st.markdown("**Ahoj! Říkám si, co tě zajímá z těch dat?** Vyber, co nejlépe odpovídá tvému výzkumu:")
    st.markdown('</div>', unsafe_allow_html=True)
    
    goal = st.radio(
        "Můj výzkumný cíl:",
        [
            "🔗 Chci zjistit, jestli spolu dvě věci souvisí (korelace/vztah)",
            "⚖️ Chci porovnat skupiny (jsou mezi nimi rozdíly?)",
            "📋 Chci popsat svá data (průměry, četnosti)",
        ],
        label_visibility="collapsed"
    )
    
    if "souvisí" in goal:
        st.session_state.research_goal = "vztah"
        st.markdown('<div class="info-box">', unsafe_allow_html=True)
        st.markdown("**Perfektní volba!** 🔗 Budu hledat, jak silně spolu proměnné souvisí.")
        st.markdown("**Příklad:** *Souvisí počet hodin studia se známkami? Souvisí výška s váhou?*")
        st.markdown('</div>', unsafe_allow_html=True)
    elif "porovnat" in goal:
        st.session_state.research_goal = "porovnání"
        st.markdown('<div class="info-box">', unsafe_allow_html=True)
        st.markdown("**Skvělé!** ⚖️ Zjistím, jestli se skupiny statisticky liší.")
        st.markdown("**Příklad:** *Mají muži a ženy jiné platy? Liší se známky mezi obory?*")
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        st.session_state.research_goal = "popis"
        st.markdown('<div class="info-box">', unsafe_allow_html=True)
        st.markdown("**Výborně!** 📋 Připravím přehledné statistiky tvých dat.")
        st.markdown("**Příklad:** *Jaký je průměrný věk? Kolik lidí vybralo každou odpověď?*")
        st.markdown('</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("⬅️ Zpět na data", use_container_width=True):
            st.session_state.step = 1
            st.rerun()
    with col2:
        if st.button("➡️ Ukaž mi doporučení", use_container_width=True):
            recommendations = recommend_analysis(df, types, st.session_state.research_goal)
            st.session_state.recommended_analysis = recommendations
            st.session_state.step = 3
            st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# ========================================
# KROK 3: DOPORUČENÍ A ANALÝZA
# ========================================

elif st.session_state.step == 3:
    df = st.session_state.df
    types, _ = detect_var_types(df)
    recommendations = st.session_state.recommended_analysis
    
    st.markdown('<div class="step-card">', unsafe_allow_html=True)
    st.markdown("## 🎓 Krok 3: Doporučení a analýza")
    
    st.markdown('<div class="chat-bubble">', unsafe_allow_html=True)
    st.markdown(f"**Na základě tvých dat a cíle '{st.session_state.research_goal}' ti doporučuji tyto analýzy:**")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Zobrazení doporučení
    for i, rec in enumerate(recommendations):
        with st.expander(f"{rec['icon']} **{rec['name']}** — {rec['reason']}", expanded=(i==0)):
            st.markdown(f"**Vhodné pro:** {rec['suitable_for']}")
            st.markdown(f"**Příklad použití:** {rec['example']}")
            
            if rec['type'] and st.button(f"✨ Použít tuto analýzu", key=f"btn_{i}", use_container_width=True):
                st.session_state.selected_analysis = rec
                st.session_state.step = 4
                st.rerun()
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("⬅️ Zpět na cíl", use_container_width=True):
            st.session_state.step = 2
            st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# ========================================
# KROK 4: PROVEDENÍ ANALÝZY
# ========================================

elif st.session_state.step == 4:
    df = st.session_state.df
    types, _ = detect_var_types(df)
    analysis = st.session_state.selected_analysis
    
    st.markdown('<div class="step-card">', unsafe_allow_html=True)
    st.markdown(f"## {analysis['icon']} {analysis['name']}")
    
    result_text = ""
    fig = None
    table_data = None
    
    # KORELACE
    if analysis['type'] == 'correlation':
        st.markdown("### Vyber proměnné pro korelaci:")
        col1, col2, col3 = st.columns(3)
        with col1:
            x = st.selectbox("Proměnná X:", analysis['variables']['x'])
        with col2:
            y = st.selectbox("Proměnná Y:", [v for v in analysis['variables']['y'] if v != x])
        with col3:
            method = st.radio("Metoda:", ["Pearson", "Spearman"])
        
        warnings, tips = check_assumptions(df, 'correlation', x, y)
        
        if warnings:
            st.markdown('<div class="warning-box">', unsafe_allow_html=True)
            for w in warnings:
                st.markdown(w)
            st.markdown('</div>', unsafe_allow_html=True)
        
        if tips:
            st.markdown('<div class="success-box">', unsafe_allow_html=True)
            for t in tips:
                st.markdown(t)
            st.markdown('</div>', unsafe_allow_html=True)
        
        if st.button("🚀 Spustit analýzu", use_container_width=True):
            sx = df[x].dropna()
            sy = df[y].dropna()
            common_idx = sx.index.intersection(sy.index)
            sx = sx[common_idx]
            sy = sy[common_idx]
            
            if method == "Pearson":
                r, p = stats.pearsonr(sx, sy)
            else:
                r, p = stats.spearmanr(sx, sy)
            
            # Interpretace
            strength = "velmi slabý"
            if abs(r) >= 0.7:
                strength = "silný"
            elif abs(r) >= 0.5:
                strength = "středně silný"
            elif abs(r) >= 0.3:
                strength = "slabý"
            
            direction = "pozitivní" if r > 0 else "negativní"
            significant = "**statisticky významný**" if p < 0.05 else "**není statisticky významný**"
            
            result_text = f"""
### 📊 Výsledky korelační analýzy

**Metoda:** {method}ova korelace  
**Počet pozorování:** {len(sx)}  
**Korelační koeficient (r):** {r:.3f}  
**P-hodnota:** {p:.4f}  

### 💡 Co to znamená pro tvou práci:

Mezi proměnnými "{x}" a "{y}" byl nalezen **{strength} {direction} vztah** (r = {r:.3f}).  
Tento vztah je {significant} (p = {p:.4f}).

**Pro závěrečnou práci můžeš napsat:**  
"Byla provedena {method}ova korelační analýza pro zjištění vztahu mezi {x} a {y} (n = {len(sx)}). 
Výsledky ukázaly {strength} {direction} vztah (r = {r:.3f}, p = {p:.4f}), 
který {"je" if p < 0.05 else "není"} statisticky významný na hladině α = 0,05."
            """
            
            # Graf
            fig = px.scatter(df, x=x, y=y, trendline="ols", 
                           title=f"Rozptylový graf: {x} vs {y}",
                           labels={x: x, y: y})
            fig.update_layout(
                font=dict(size=14),
                plot_bgcolor='white',
                paper_bgcolor='white',
                title_font_size=18
            )
            
            # Tabulka
            table_data = pd.DataFrame({
                'Metrika': ['Korelační koeficient (r)', 'P-hodnota', 'Počet pozorování', 'Síla vztahu'],
                'Hodnota': [f"{r:.3f}", f"{p:.4f}", str(len(sx)), strength]
            })
            
            st.markdown(result_text)
            st.plotly_chart(fig, use_container_width=True)
            st.markdown("### 📋 Tabulka výsledků (ready for copy-paste)")
            st.dataframe(table_data, use_container_width=True)
    
    # T-TEST
    elif analysis['type'] == 'ttest':
        st.markdown("### Vyber proměnné pro t-test:")
        col1, col2 = st.columns(2)
        with col1:
            group = st.selectbox("Skupinová proměnná (2 kategorie):", analysis['variables']['group'])
        with col2:
            outcome = st.selectbox("Měřená proměnná (číslo):", analysis['variables']['outcome'])
        
        warnings, tips = check_assumptions(df, 'ttest', group, outcome)
        
        if warnings:
            st.markdown('<div class="warning-box">', unsafe_allow_html=True)
            for w in warnings:
                st.markdown(w)
            st.markdown('</div>', unsafe_allow_html=True)
        
        if tips:
            st.markdown('<div class="success-box">', unsafe_allow_html=True)
            for t in tips:
                st.markdown(t)
            st.markdown('</div>', unsafe_allow_html=True)
        
        if st.button("🚀 Spustit analýzu", use_container_width=True):
            tmp = df[[group, outcome]].dropna()
            groups = tmp[group].unique()
            
            g1 = tmp[tmp[group] == groups[0]][outcome]
            g2 = tmp[tmp[group] == groups[1]][outcome]
            
            t, p = stats.ttest_ind(g1, g2)
            
            # Cohen's d
            n1, n2 = len(g1), len(g2)
            s1, s2 = np.var(g1, ddof=1), np.var(g2, ddof=1)
            sp = np.sqrt(((n1-1)*s1 + (n2-1)*s2) / (n1+n2-2))
            d = (np.mean(g1) - np.mean(g2)) / sp if sp != 0 else 0
            
            effect_size = "malý"
            if abs(d) >= 0.8:
                effect_size = "velký"
            elif abs(d) >= 0.5:
                effect_size = "střední"
            
            significant = "**statisticky významný**" if p < 0.05 else "**není statisticky významný**"
            
            result_text = f"""
### 📊 Výsledky t-testu

**Porovnání:** {groups[0]} vs {groups[1]}  
**Skupina 1 ({groups[0]}):** průměr = {np.mean(g1):.2f}, n = {n1}  
**Skupina 2 ({groups[1]}):** průměr = {np.mean(g2):.2f}, n = {n2}  
**T-statistika:** {t:.3f}  
**P-hodnota:** {p:.4f}  
**Cohenovo d:** {d:.3f} ({effect_size} efekt)

### 💡 Co to znamená pro tvou práci:

Rozdíl v průměrech proměnné "{outcome}" mezi skupinami "{groups[0]}" a "{groups[1]}" je {significant} (p = {p:.4f}).
Velikost efektu je {effect_size} (d = {d:.3f}).

**Pro závěrečnou práci můžeš napsat:**  
"Pro porovnání průměrů proměnné {outcome} mezi skupinami {group} byl použit dvouvýběrový t-test (n₁ = {n1}, n₂ = {n2}). 
Průměr skupiny {groups[0]} byl {np.mean(g1):.2f}, průměr skupiny {groups[1]} byl {np.mean(g2):.2f}. 
Rozdíl {"je" if p < 0.05 else "není"} statisticky významný (t = {t:.2f}, p = {p:.4f}, d = {d:.2f})."
            """
            
            # Graf
            plot_df = pd.DataFrame({
                group: list(g1.index.map(lambda x: groups[0])) + list(g2.index.map(lambda x: groups[1])),
                outcome: list(g1) + list(g2)
            })
            fig = px.box(plot_df, x=group, y=outcome, title=f"Porovnání: {outcome} podle {group}",
                        color=group)
            fig.update_layout(showlegend=False, font=dict(size=14))
            
            # Tabulka
            table_data = pd.DataFrame({
                'Skupina': [groups[0], groups[1]],
                'Průměr': [f"{np.mean(g1):.2f}", f"{np.mean(g2):.2f}"],
                'Směrodatná odchylka': [f"{np.std(g1, ddof=1):.2f}", f"{np.std(g2, ddof=1):.2f}"],
                'n': [n1, n2]
            })
            
            st.markdown(result_text)
            st.plotly_chart(fig, use_container_width=True)
            st.markdown("### 📋 Tabulka deskriptivních statistik")
            st.dataframe(table_data, use_container_width=True)
    
    # ANOVA
    elif analysis['type'] == 'anova':
        st.markdown("### Vyber proměnné pro ANOVA:")
        col1, col2 = st.columns(2)
        with col1:
            group = st.selectbox("Skupinová proměnná (3+ kategorií):", analysis['variables']['group'])
        with col2:
            outcome = st.selectbox("Měřená proměnná (číslo):", analysis['variables']['outcome'])
        
        if st.button("🚀 Spustit analýzu", use_container_width=True):
            tmp = df[[group, outcome]].dropna()
            formula = f"{outcome} ~ C({group})"
            model = ols(formula, data=tmp).fit()
            anova_table = sm.stats.anova_lm(model, typ=2)
            
            F = anova_table.loc[f'C({group})', 'F']
            p = anova_table.loc[f'C({group})', 'PR(>F)']
            
            # Eta squared
            ss_effect = anova_table.loc[f'C({group})', 'sum_sq']
            ss_resid = anova_table.loc['Residual', 'sum_sq']
            eta2 = ss_effect / (ss_effect + ss_resid)
            
            effect_size = "malý"
            if eta2 >= 0.14:
                effect_size = "velký"
            elif eta2 >= 0.06:
                effect_size = "střední"
            
            significant = "**statisticky významný**" if p < 0.05 else "**není statisticky významný**"
            
            # Průměry skupin
            group_means = tmp.groupby(group)[outcome].agg(['mean', 'std', 'count'])
            
            result_text = f"""
### 📊 Výsledky ANOVA

**F-statistika:** {F:.3f}  
**P-hodnota:** {p:.4f}  
**Eta-squared (η²):** {eta2:.3f} ({effect_size} efekt)  
**Počet skupin:** {tmp[group].nunique()}  
**Celkový počet pozorování:** {len(tmp)}

### 💡 Co to znamená pro tvou práci:

Rozdíly v průměrech proměnné "{outcome}" mezi skupinami "{group}" jsou {significant} (p = {p:.4f}).
Velikost efektu je {effect_size} (η² = {eta2:.3f}).

**Pro závěrečnou práci můžeš napsat:**  
"Pro porovnání průměrů proměnné {outcome} mezi skupinami {group} byla použita jednofaktorová ANOVA (n = {len(tmp)}). 
Rozdíly mezi skupinami {"jsou" if p < 0.05 else "nejsou"} statisticky významné (F = {F:.2f}, p = {p:.4f}, η² = {eta2:.2f})."
            """
            
            # Graf
            fig = px.box(tmp, x=group, y=outcome, title=f"Porovnání: {outcome} podle {group}",
                        color=group)
            fig.update_layout(showlegend=False, font=dict(size=14))
            
            # Tabulka
            table_data = group_means.reset_index()
            table_data.columns = ['Skupina', 'Průměr', 'Směrodatná odchylka', 'n']
            table_data['Průměr'] = table_data['Průměr'].round(2)
            table_data['Směrodatná odchylka'] = table_data['Směrodatná odchylka'].round(2)
            
            st.markdown(result_text)
            st.plotly_chart(fig, use_container_width=True)
            st.markdown("### 📋 Tabulka deskriptivních statistik")
            st.dataframe(table_data, use_container_width=True)
    
    # CHÍ-KVADRÁT
    elif analysis['type'] == 'chi2':
        st.markdown("### Vyber proměnné pro chí-kvadrát test:")
        col1, col2 = st.columns(2)
        with col1:
            var1 = st.selectbox("Proměnná 1:", analysis['variables']['var1'])
        with col2:
            var2 = st.selectbox("Proměnná 2:", [v for v in analysis['variables']['var2'] if v != var1])
        
        if st.button("🚀 Spustit analýzu", use_container_width=True):
            tmp = df[[var1, var2]].dropna()
            ctab = pd.crosstab(tmp[var1], tmp[var2])
            chi2, p, dof, expected = stats.chi2_contingency(ctab)
            
            # Cramér's V
            n = ctab.sum().sum()
            r, c = ctab.shape
            v = np.sqrt(chi2 / (n * (min(r-1, c-1))))
            
            effect_size = "slabá"
            if v >= 0.5:
                effect_size = "silná"
            elif v >= 0.3:
                effect_size = "střední"
            
            significant = "**statisticky významná**" if p < 0.05 else "**není statisticky významná**"
            
            result_text = f"""
### 📊 Výsledky chí-kvadrát testu

**Chí-kvadrát statistika:** {chi2:.3f}  
**P-hodnota:** {p:.4f}  
**Stupně volnosti:** {dof}  
**Cramérovo V:** {v:.3f} ({effect_size} asociace)  
**Počet pozorování:** {n}

### 💡 Co to znamená pro tvou práci:

Asociace mezi proměnnými "{var1}" a "{var2}" je {significant} (p = {p:.4f}).
Síla asociace je {effect_size} (V = {v:.3f}).

**Pro závěrečnou práci můžeš napsat:**  
"Pro testování nezávislosti mezi proměnnými {var1} a {var2} byl použit chí-kvadrát test (n = {n}). 
Asociace mezi proměnnými {"je" if p < 0.05 else "není"} statisticky významná (χ² = {chi2:.2f}, df = {dof}, p = {p:.4f}, V = {v:.2f})."
            """
            
            # Graf
            fig = px.bar(ctab.reset_index().melt(id_vars=var1), 
                        x=var1, y='value', color=var2, barmode='group',
                        title=f"Kontingenční tabulka: {var1} × {var2}")
            fig.update_layout(font=dict(size=14))
            
            st.markdown(result_text)
            st.plotly_chart(fig, use_container_width=True)
            st.markdown("### 📋 Kontingenční tabulka")
            st.dataframe(ctab, use_container_width=True)
    
    # LINEÁRNÍ REGRESE
    elif analysis['type'] == 'regression':
        st.markdown("### Vyber proměnné pro regresi:")
        col1, col2 = st.columns(2)
        with col1:
            x_var = st.selectbox("Nezávislá proměnná (X, prediktor):", analysis['variables']['x'])
        with col2:
            y_var = st.selectbox("Závislá proměnná (Y, výstup):", [v for v in analysis['variables']['y'] if v != x_var])
        
        if st.button("🚀 Spustit analýzu", use_container_width=True):
            sx = df[x_var].dropna()
            sy = df[y_var].dropna()
            common_idx = sx.index.intersection(sy.index)
            sx = sx[common_idx]
            sy = sy[common_idx]
            
            X = sm.add_constant(sx)
            model = sm.OLS(sy, X).fit()
            
            r2 = model.rsquared
            adj_r2 = model.rsquared_adj
            f_stat = model.fvalue
            f_p = model.f_pvalue
            coef = model.params[x_var]
            coef_p = model.pvalues[x_var]
            intercept = model.params['const']
            
            significant = "**statisticky významný**" if f_p < 0.05 else "**není statisticky významný**"
            
            result_text = f"""
### 📊 Výsledky lineární regrese

**Model:** {y_var} = {intercept:.3f} + {coef:.3f} × {x_var}  
**R² (procento vysvětlené variability):** {r2:.3f} ({r2*100:.1f}%)  
**Adjustované R²:** {adj_r2:.3f}  
**F-statistika:** {f_stat:.3f}, p = {f_p:.4f}  
**Koeficient pro {x_var}:** {coef:.3f} (p = {coef_p:.4f})  
**Počet pozorování:** {len(sx)}

### 💡 Co to znamená pro tvou práci:

Model je {significant} (p = {f_p:.4f}).  
Proměnná "{x_var}" vysvětluje {r2*100:.1f}% variability proměnné "{y_var}".  
Při zvýšení {x_var} o 1 jednotku se {y_var} {"zvýší" if coef > 0 else "sníží"} průměrně o {abs(coef):.3f}.

**Pro závěrečnou práci můžeš napsat:**  
"Byla provedena jednoduchá lineární regrese pro predikci {y_var} na základě {x_var} (n = {len(sx)}). 
Model {"je" if f_p < 0.05 else "není"} statisticky významný (F = {f_stat:.2f}, p = {f_p:.4f}) 
a vysvětluje {r2*100:.1f}% variability závislé proměnné (R² = {r2:.3f})."
            """
            
            # Graf
            fig = px.scatter(df, x=x_var, y=y_var, trendline="ols",
                           title=f"Lineární regrese: {y_var} ~ {x_var}")
            fig.update_layout(font=dict(size=14))
            
            # Tabulka
            table_data = pd.DataFrame({
                'Metrika': ['R²', 'Adjustované R²', 'F-statistika', 'P-hodnota', 'Koeficient', 'Intercept'],
                'Hodnota': [f"{r2:.3f}", f"{adj_r2:.3f}", f"{f_stat:.2f}", f"{f_p:.4f}", f"{coef:.3f}", f"{intercept:.3f}"]
            })
            
            st.markdown(result_text)
            st.plotly_chart(fig, use_container_width=True)
            st.markdown("### 📋 Tabulka výsledků")
            st.dataframe(table_data, use_container_width=True)
    
    # DESKRIPTIVNÍ STATISTIKA
    elif analysis['type'] == 'descriptive':
        st.markdown("### Vyber proměnné pro popis:")
        selected_cols = st.multiselect("Které proměnné chceš popsat?", df.columns.tolist(), default=df.columns.tolist()[:3])
        
        if st.button("🚀 Zobrazit statistiky", use_container_width=True):
            num_cols = [c for c in selected_cols if pd.api.types.is_numeric_dtype(df[c])]
            cat_cols = [c for c in selected_cols if not pd.api.types.is_numeric_dtype(df[c])]
            
            st.markdown("### 📊 Numerické proměnné")
            if num_cols:
                desc = df[num_cols].describe().T
                desc['chybějící'] = df[num_cols].isnull().sum()
                st.dataframe(desc.round(2), use_container_width=True)
                
                for col in num_cols:
                    fig = px.histogram(df, x=col, title=f"Distribuce: {col}")
                    st.plotly_chart(fig, use_container_width=True)
            
            st.markdown("### 🏷️ Kategorické proměnné")
            if cat_cols:
                for col in cat_cols:
                    freq = df[col].value_counts()
                    st.markdown(f"**{col}:**")
                    st.dataframe(freq, use_container_width=True)
                    
                    fig = px.bar(x=freq.index, y=freq.values, title=f"Četnosti: {col}",
                               labels={'x': col, 'y': 'Počet'})
                    st.plotly_chart(fig, use_container_width=True)
    
    # EXPORT
    if result_text:
        st.markdown("---")
        st.markdown("### 💾 Export výsledků")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📄 Stáhnout jako Word", use_container_width=True):
                doc = Document()
                
                # Titulek
                title = doc.add_heading(f'QuantBuddy — {analysis["name"]}', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Metadata
                doc.add_paragraph(f"Datum: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                doc.add_paragraph(f"Počet pozorování: {len(df)}")
                
                # Výsledky
                doc.add_heading('Výsledky analýzy', level=1)
                doc.add_paragraph(result_text.replace('###', '').replace('**', ''))
                
                # Graf
                if fig:
                    doc.add_heading('Vizualizace', level=1)
                    img_bytes = fig.to_image(format="png", width=1200, height=800)
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        tmp.write(img_bytes)
                        tmp.flush()
                        doc.add_picture(tmp.name, width=Inches(6))
                
                # Uložení
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                st.download_button(
                    label="⬇️ Stáhnout Word dokument",
                    data=bio,
                    file_name=f"quantbuddy_{analysis['type']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with col2:
            if st.button("🔄 Spustit novou analýzu", use_container_width=True):
                st.session_state.step = 2
                st.rerun()
    
    if st.button("⬅️ Zpět na doporučení", use_container_width=True):
        st.session_state.step = 3
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# ========================================
# FOOTER
# ========================================

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: gray; font-size: 0.9rem; padding: 2rem 0;">
    <p>🤖 <strong>QuantBuddy AI v3.0</strong></p>
    <p>Tvůj inteligentní asistent pro kvantitativní výzkum</p>
    <p style="font-size: 0.8rem; margin-top: 1rem;">
        💡 <em>Tip: Výsledky jsou připravené pro použití v závěrečné práci, 
        ale vždy je dobré konzultovat s vedoucím!</em>
    </p>
</div>
""", unsafe_allow_html=True)
